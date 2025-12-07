#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("<span style='color:red'>❌ Error: La librería 'python-docx' no está instalada.</span>")
    sys.exit(1)

def generar_manual():
    print(">>> Generando 'Media Manager Pro: Guía Maestra' (Edición Extendida)...")
    
    output_folder = "/app/datos"
    filename = "Media_Manager_Pro_Manual.docx"
    output_path = os.path.join(output_folder, filename)

    if not os.path.exists(output_folder):
        try: os.makedirs(output_folder)
        except: output_path = filename 

    try:
        doc = Document()

        # ==========================================
        # 1. ESTILOS VISUALES (Editorial Técnica)
        # ==========================================
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)
        style_normal.paragraph_format.space_after = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.15

        # Títulos
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            style.font.name = 'Segoe UI'
            style.font.bold = True
            if i == 1: 
                style.font.size = Pt(22); style.font.color.rgb = RGBColor(0, 51, 102)
                style.paragraph_format.space_before = Pt(24); style.paragraph_format.page_break_before = True
            if i == 2: 
                style.font.size = Pt(16); style.font.color.rgb = RGBColor(45, 112, 184)
                style.paragraph_format.space_before = Pt(18); style.paragraph_format.page_break_before = False
            if i == 3:
                style.font.size = Pt(13); style.font.color.rgb = RGBColor(80, 80, 80)

        # Helpers Visuales
        def add_code_block(code_text):
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F4F8"/>'.format(nsdecls('w'))) # Gris azulado muy suave
            cell._tc.get_or_add_tcPr().append(shading_elm)
            p = cell.paragraphs[0]
            p.style = 'No Spacing'
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 30, 30)
            p.paragraph_format.left_indent = Pt(5)
            p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
            doc.add_paragraph() 

        def add_info_box(title, text):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            # Fondo Amarillo muy suave (#FFFFE0)
            shading_elm = parse_xml(r'<w:shd {} w:fill="FFF9C4"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            p = cell.paragraphs[0]
            p.add_run(f"💡 {title}: ").bold = True
            p.add_run(text)
            doc.add_paragraph()

        def add_styled_table(headers, rows):
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Medium Shading 1 Accent 1'
            hdr = table.rows[0].cells
            for i, txt in enumerate(headers): hdr[i].text = txt
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row): cells[i].text = str(val)
            doc.add_paragraph()

        # ==========================================
        # PORTADA
        # ==========================================
        doc.styles['Heading 1'].paragraph_format.page_break_before = False
        for _ in range(5): doc.add_paragraph()
        
        t = doc.add_heading('MEDIA MANAGER PRO', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].font.size = Pt(40); t.runs[0].font.color.rgb = RGBColor(0, 32, 96)

        sub = doc.add_paragraph('Manual de Ingeniería DevOps para Unraid')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(16); sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        for _ in range(8): doc.add_paragraph()
        
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Esta guía está diseñada para llevarte de principiante a experto en la gestión de código y contenedores.").italic = True
        
        doc.add_page_break()
        doc.styles['Heading 1'].paragraph_format.page_break_before = True

        # ==========================================
        # CAPÍTULO 1: FUNDAMENTOS
        # ==========================================
        doc.add_heading('1. CONCEPTOS FUNDAMENTALES', level=1)
        doc.add_paragraph("Antes de escribir comandos, es vital entender qué estamos construyendo. Este sistema se basa en tres pilares que trabajan juntos:")

        doc.add_heading('El Cerebro: Git (Control de Versiones)', level=2)
        doc.add_paragraph("Imagina que Git es una máquina del tiempo para tus archivos. Cada vez que haces un 'commit', guardas una foto exacta de cómo estaba tu proyecto en ese momento. Si rompes algo mañana, puedes volver a la foto de hoy.")
        add_info_box("Ramas (Branches)", "Trabajamos con dos líneas temporales paralelas:\n• Rama MAIN: Es el producto final. Como lo que ves en una tienda. No se toca, solo se admira.\n• Rama DEV: Es el taller sucio. Aquí cortas, pegas y pruebas. Si algo explota aquí, no pasa nada.")

        doc.add_heading('El Cuerpo: Docker (Contenedores)', level=2)
        doc.add_paragraph("Un Docker no es una máquina virtual, es más como un 'tupper' hermético. Dentro metes tu código y todas las librerías que necesita (Python, Flask, etc.).")
        doc.add_paragraph("La magia es que este tupper funciona IGUAL en tu ordenador, en mi servidor o en la nube. Se acabó el 'en mi máquina funcionaba'.")
        
        doc.add_heading('El Almacén: GHCR (GitHub Container Registry)', level=2)
        doc.add_paragraph("Es como un Dropbox, pero solo para tus 'tuppers' (imágenes Docker). Una vez cocinas tu imagen en tu servidor, la subes aquí para tener una copia de seguridad perfecta en la nube.")

        # ==========================================
        # CAPÍTULO 2: ARQUITECTURA
        # ==========================================
        doc.add_heading('2. ARQUITECTURA EN UNRAID', level=1)
        doc.add_paragraph("En Unraid, necesitamos mapear (conectar) carpetas de tu disco duro real (Host) hacia dentro del contenedor.")

        doc.add_heading('Mapa de Directorios (Volúmenes)', level=2)
        table_data = [['Ruta Unraid (Host)', 'Ruta Docker (Interna)', 'Explicación para Humanos']]
        table_data.append(['.../scripts/media-manager-working/', '/app/scripts', 'CÓDIGO FUENTE. Aquí es donde trabajas con Git.'])
        table_data.append(['.../media-manager-pro/', '/app/datos', 'PRODUCCIÓN. Base de datos real. ¡Cuidado!'])
        table_data.append(['.../media-manager-dev/', '/app/datos', 'PRUEBAS. Base de datos de juguete. Rompe sin miedo.'])
        add_styled_table(table_data[0], table_data[1:])

        doc.add_heading('Estrategia de Doble Contenedor', level=2)
        doc.add_paragraph("Para no dejar de tener servicio mientras programas, usamos dos contenedores simultáneos:")
        doc.add_paragraph("1. Media-Manager-PRO (Puerto 5000): Siempre encendido. Solo se actualiza cuando estamos seguros de que todo funciona.")
        doc.add_paragraph("2. Media-Manager-DEV (Puerto 5001): Tu zona de guerra. Aquí reinicias constantemente para ver tus cambios.")

        # ==========================================
        # CAPÍTULO 3: SETUP INICIAL
        # ==========================================
        doc.add_heading('3. INICIALIZACIÓN Y RESET (Primer Setup)', level=1)
        
        warn = doc.add_paragraph()
        run_warn = warn.add_run("¡STOP! 🛑 Lee esto antes de seguir.")
        run_warn.bold = True; run_warn.font.color.rgb = RGBColor(200, 0, 0)
        doc.add_paragraph("Este capítulo contiene instrucciones destructivas. Solo debes ejecutar esto si estás instalando el proyecto por primera vez o si quieres borrar toda la historia de Git y empezar de cero absoluto.")

        doc.add_heading('Paso 1: Limpieza y Creación del Repositorio', level=2)
        doc.add_paragraph("Vamos a decirle a tu carpeta: 'Olvida todo tu pasado y empieza una vida nueva'.")
        add_code_block("cd /mnt/user/appdata/scripts/media-manager-working/\nrm -rf .git  # Borramos el cerebro de Git anterior\ngit init       # Creamos un cerebro nuevo y vacío\ngit add .      # Le decimos: 'Mira todos estos archivos'\ngit commit -m \"Initial commit: Nace el proyecto\"")

        doc.add_heading('Paso 2: Conexión con la Nube', level=2)
        doc.add_paragraph("Ahora tu código vive en tu disco duro, pero si el disco muere, el código muere. Vamos a subirlo a GitHub.")
        add_code_block("git remote add origin https://github.com/jvdivx/media-manager-unraid.git\ngit branch -M main\ngit push -u --force origin main")
        add_info_box("Truco de Autenticación", "Cuando la terminal te pida 'Password', NO escribas tu contraseña de GitHub. Pega tu Token personal que empieza por 'ghp_'.")

        doc.add_heading('Paso 3: Crear la Rama de Pruebas', level=2)
        doc.add_paragraph("Nunca programamos en la rama principal. Creamos una copia llamada 'dev'.")
        add_code_block("git checkout -b dev\ngit push -u origin dev")

        # ==========================================
        # CAPÍTULO 4: FLUJO DIARIO
        # ==========================================
        doc.add_heading('4. TU DÍA A DÍA (Ciclo de Desarrollo)', level=1)
        doc.add_paragraph("Esto es lo que harás el 99% de las veces. Memoriza este ciclo:")
        doc.add_paragraph("PROGRAMAR -> GUARDAR (Git) -> CONSTRUIR (Docker) -> PROBAR")

        doc.add_heading('A. Preparar el Terreno', level=2)
        doc.add_paragraph("Antes de empezar, asegúrate de estar en la rama correcta y tener lo último.")
        add_code_block("cd /mnt/user/appdata/scripts/media-manager-working/\ngit checkout dev\ngit pull origin dev")

        doc.add_heading('B. Guardar tus Cambios (Git Push)', level=2)
        doc.add_paragraph("Has editado un script y funciona en tu PC. Es hora de guardarlo en la nube.")
        add_code_block("git status    # Mira qué archivos están en rojo (modificados)\ngit add .     # Ponlos todos en verde (listos para guardar)\ngit commit -m \"Añado función de reporte PDF\"  # Ponle etiqueta\ngit push origin dev  # Envíalo a GitHub")

        doc.add_heading('C. Actualizar el Contenedor (Docker Build)', level=2)
        doc.add_paragraph("Tu código está guardado, pero el contenedor 'Media-Manager-DEV' sigue ejecutando la versión vieja. Necesitamos 'reconstruirlo'.")
        add_code_block("docker build -t ghcr.io/jvdivx/media-manager-pro:dev .\ndocker push ghcr.io/jvdivx/media-manager-pro:dev")
        add_info_box("¿Y ahora qué?", "Ve a la pestaña Docker de Unraid y reinicia el contenedor 'Media-Manager-DEV'. ¡Tus cambios ya estarán activos en el puerto 5001!")

        # ==========================================
        # CAPÍTULO 5: PRODUCCIÓN
        # ==========================================
        doc.add_heading('5. LLEVAR A PRODUCCIÓN (Merge)', level=1)
        doc.add_paragraph("Solo haz esto cuando estés 100% seguro de que tu código en DEV funciona perfecto. Vamos a mover los cambios de la 'zona sucia' a la 'zona limpia'.")

        doc.add_heading('El Baile de las Ramas', level=2)
        doc.add_paragraph("1. Nos cambiamos a la rama principal (main).")
        doc.add_paragraph("2. Absorbemos (merge) todo lo nuevo de dev.")
        doc.add_paragraph("3. Subimos a la nube.")
        doc.add_paragraph("4. Volvemos corriendo a dev para seguir trabajando.")
        add_code_block("git checkout main\ngit merge dev\ngit push origin main\n\n# Generar la imagen Oficial\ndocker build -t ghcr.io/jvdivx/media-manager-pro:latest .\ndocker push ghcr.io/jvdivx/media-manager-pro:latest\n\ngit checkout dev  # ¡Importante volver!")

        # ==========================================
        # CAPÍTULO 6: MANTENIMIENTO
        # ==========================================
        doc.add_heading('6. MANTENIMIENTO Y LIMPIEZA', level=1)
        
        doc.add_heading('Limpieza de Disco Docker', level=2)
        doc.add_paragraph("Cada vez que haces un 'build', Docker crea capas nuevas. Las viejas se quedan ocupando espacio como 'basura'.")
        doc.add_paragraph("Si Unraid te avisa de que el 'Docker vDisk' está lleno, ejecuta esto:")
        add_code_block("docker image prune -f  # Borra lo superficial\ndocker system prune -a # Borra TODO lo que no se esté usando (nuclear)")

        doc.add_heading('Añadir nuevas Librerías Python', level=2)
        doc.add_paragraph("Si tu script necesita una librería nueva (ej: 'pandas'), no basta con instalarla. Debes añadirla al archivo 'requirements.txt' y volver a hacer un 'docker build'. Así quedará instalada para siempre.")

        # ==========================================
        # CAPÍTULO 7: TROUBLESHOOTING
        # ==========================================
        doc.add_heading('7. SOLUCIÓN DE PROBLEMAS COMUNES', level=1)
        
        probs = [
            ['Error', 'Qué significa', 'Solución'],
            ['Conflict Merge', 'Tocaste la misma línea en Main y Dev.', 'Abre el archivo, borra las marcas <<<< HEAD y haz commit.'],
            ['Manifest unknown', 'Docker no encuentra la imagen en la nube.', 'Haz docker build y docker push primero.'],
            ['Permission denied', 'El script no puede escribir en disco.', 'Ejecuta el script de "Reparar Permisos" (Opción 2).']
        ]
        add_styled_table(probs[0], probs[1:])

        # ==========================================
        # CAPÍTULO 8: CHULETA
        # ==========================================
        doc.add_heading('8. CHULETA RÁPIDA (CHEAT SHEET)', level=1)
        doc.add_paragraph("Imprime esta página y pégala en tu pared.")
        
        add_code_block("# GIT BÁSICO\ngit status              -> ¿Qué he tocado?\ngit add .               -> Preparar todo\ngit commit -m \"msg\"     -> Guardar foto\ngit push origin dev     -> Subir a nube\n\n# DOCKER BÁSICO\ndocker ps               -> ¿Qué corre?\ndocker logs -f [id]     -> Ver consola en vivo\ndocker exec -it [id] bash -> Entrar dentro")

        # FINAL
        doc.save(output_path)
        print(f"✅ GUÍA MAESTRA GENERADA: {output_path}")
        print("⬇️ Descárgalo desde el panel de control.")

    except Exception as e:
        print(f"<span style='color:red'>❌ Error generando el manual: {e}</span>")

if __name__ == "__main__":
    generar_manual()
